from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/Guion_explicacion_proyecto_horarios.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.name = "Calibri"
    r2.font.size = Pt(10)
    doc.add_paragraph()


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "F2F4F7")
        set_cell_text(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Guion para explicar el proyecto")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Generador de Horarios Escolares con Algoritmo Genetico")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(31, 77, 120)
    subtitle.paragraph_format.space_after = Pt(14)

    add_note(
        doc,
        "Idea principal para iniciar la exposicion",
        "Mi proyecto resuelve el problema de acomodar profesores, grupos y materias en una semana escolar, evitando choques de horario y reduciendo huecos muertos mediante un algoritmo evolutivo.",
    )

    add_heading(doc, "1. Contexto del proyecto", 1)
    doc.add_paragraph(
        "El proyecto corresponde al Generador de Horarios Escolares solicitado para la evaluacion practica de Computo Evolutivo. "
        "El objetivo es construir una aplicacion web funcional que reciba datos escolares, genere un horario y muestre los resultados de forma visual."
    )
    add_bullets(
        doc,
        [
            "Problema: acomodar profesores, grupos y materias en dias y horas disponibles.",
            "Reglas estrictas: un profesor no puede estar en dos clases al mismo tiempo; un grupo tampoco; y todas las horas semanales deben asignarse.",
            "Meta: obtener un horario valido y comodo, penalizando choques y huecos muertos.",
        ],
    )

    add_heading(doc, "2. Que hace mi sistema", 1)
    doc.add_paragraph(
        "La aplicacion permite cargar o registrar la informacion academica, configurar los horarios habiles y ejecutar un algoritmo genetico para generar una propuesta de horario. "
        "La salida se muestra en una cuadricula tipo calendario con filtros por grupo y profesor."
    )
    add_table(
        doc,
        ["Modulo", "Funcion"],
        [
            ["Carga de datos", "Importa profesores, grupos, materias y horas semanales desde CSV, Excel o PDF de horario."],
            ["Configuracion", "Permite elegir dias laborables, hora de inicio y hora final."],
            ["Algoritmo genetico", "Genera soluciones candidatas, evalua choques, aplica cruza, mutacion y reparacion."],
            ["Visualizacion", "Muestra calendario, filtros, metricas y grafica de convergencia del fitness."],
            ["Validacion", "Reporta choques de profesor, choques de grupo, horas sin asignar, total de reglas rotas y mejor fitness."],
        ],
        [1.55, 4.95],
    )

    add_heading(doc, "3. Arquitectura de la aplicacion", 1)
    doc.add_paragraph(
        "El sistema esta dividido en frontend, backend y base de datos. Esta separacion permite que la interfaz se enfoque en la experiencia del usuario y el backend en la logica de importacion, validacion y optimizacion."
    )
    add_table(
        doc,
        ["Parte", "Tecnologia", "Responsabilidad"],
        [
            ["Frontend", "React", "Interfaz web, formularios, filtros, calendario y metricas."],
            ["Backend", "FastAPI", "Endpoints, importacion de archivos, ejecucion del algoritmo y reglas de negocio."],
            ["Base de datos", "MongoDB", "Almacena profesores, grupos, materias, carreras, configuracion y ultimo horario generado."],
            ["Parser de PDF", "pdfplumber", "Extrae tablas de horarios en formato de cuadricula."],
        ],
        [1.25, 1.35, 3.9],
    )

    add_heading(doc, "4. Datos que necesita", 1)
    doc.add_paragraph("Para que el algoritmo pueda trabajar, cada materia debe estar relacionada con un profesor, un grupo y una cantidad de horas semanales.")
    add_bullets(
        doc,
        [
            "Profesor: docente que imparte la materia.",
            "Grupo: conjunto de alumnos que recibira la clase.",
            "Materia: clase que se debe programar.",
            "Horas semanales: numero de sesiones que deben acomodarse.",
            "Rango de horario: dias, hora de inicio y hora final disponibles.",
        ],
    )
    add_note(
        doc,
        "Punto importante",
        "Cuando el archivo es un PDF de horario ya organizado, el sistema conserva las posiciones originales como horarios preferidos para no destruir una solucion que ya era valida.",
    )

    add_heading(doc, "5. Como funciona el algoritmo genetico", 1)
    doc.add_paragraph(
        "El algoritmo genetico trabaja con una poblacion de posibles horarios. Cada individuo representa una asignacion completa de sesiones a bloques de tiempo."
    )
    add_heading(doc, "Representacion", 2)
    doc.add_paragraph(
        "La solucion se representa como un vector. Cada posicion del vector corresponde a una sesion de clase, y el valor guardado indica el bloque de horario asignado. "
        "Un bloque se calcula combinando dia y hora."
    )
    add_heading(doc, "Proceso interno", 2)
    add_numbered(
        doc,
        [
            "Expandir materias: si una materia tiene 5 horas semanales, se convierte en 5 sesiones individuales.",
            "Crear poblacion inicial: se generan varios horarios candidatos.",
            "Evaluar fitness: se calculan choques de profesor, choques de grupo, huecos y penalizaciones.",
            "Seleccionar mejores individuos: se favorecen los horarios con menor fitness.",
            "Cruzar soluciones: se combinan partes de dos horarios candidatos.",
            "Mutar: se cambian algunas sesiones aleatoriamente para explorar nuevas soluciones.",
            "Reparar: se mueven sesiones que causan choques cuando existe un bloque libre posible.",
            "Repetir por generaciones hasta obtener el mejor horario encontrado.",
        ],
    )

    add_heading(doc, "6. Funcion objetivo y metricas", 1)
    doc.add_paragraph(
        "La funcion objetivo busca minimizar el fitness. Un fitness menor significa una solucion mas conveniente. "
        "Los choques se penalizan fuertemente porque violan reglas estrictas."
    )
    add_table(
        doc,
        ["Metrica", "Que significa", "Como explicarla"],
        [
            ["Choques de Profesor", "Un profesor aparece en dos clases al mismo tiempo.", "Debe quedar en cero porque es una regla inviolable."],
            ["Choques de Grupo", "Un grupo tiene dos materias en el mismo bloque.", "Debe quedar en cero porque el grupo no puede asistir a dos clases simultaneas."],
            ["Horas Sin Asignar", "Sesiones que no se pudieron colocar.", "Debe quedar en cero para cumplir todas las horas semanales."],
            ["Mejor Fitness", "Valor numerico de la mejor solucion.", "Entre menor sea, mejor; puede no ser cero si aun hay huecos o penalizaciones suaves."],
            ["Grafica de Convergencia", "Evolucion del fitness por generaciones.", "Sirve para mostrar como el algoritmo mejora con el tiempo."],
        ],
        [1.45, 2.35, 2.7],
    )

    add_heading(doc, "7. Como se cumplen los requisitos del PDF", 1)
    add_table(
        doc,
        ["Requisito", "Cumplimiento en mi sistema"],
        [
            ["Carga de base de datos", "Permite cargar CSV, Excel o PDF; tambien permite registrar datos manualmente."],
            ["Configuracion de restricciones", "Se configuran dias laborables y rango de horas, por ejemplo de 7:00 a 19:00."],
            ["Parametros del algoritmo", "Incluye tamano de poblacion, generaciones, probabilidad de mutacion y cruza."],
            ["Boton de accion", "El boton Generar Horarios ejecuta el proceso evolutivo."],
            ["Salida visual", "Muestra una cuadricula tipo calendario y filtros por grupo o profesor."],
            ["Metricas", "Muestra reglas rotas, choques, sesiones programadas, fitness y grafica de convergencia."],
        ],
        [2.2, 4.3],
    )

    add_heading(doc, "8. Resultados que debo mostrar al profesor", 1)
    add_bullets(
        doc,
        [
            "Primero mostrar la carga del archivo base o PDF.",
            "Despues mostrar que aparecen profesores, grupos y materias cargadas.",
            "Configurar el rango de horas; para el conjunto probado se recomienda 7:00 a 19:00.",
            "Presionar Generar Horarios.",
            "Mostrar que las metricas de choques quedan en cero.",
            "Usar los filtros para enseñar el horario de un grupo y luego el de un profesor.",
            "Explicar que el fitness puede ser mayor que cero por huecos o penalizaciones suaves, pero lo critico es que las reglas estrictas queden en cero.",
        ],
    )

    add_heading(doc, "9. Guion corto para exponer", 1)
    add_note(
        doc,
        "Inicio",
        "Este proyecto aplica computo evolutivo a un problema real de optimizacion: la generacion automatica de horarios escolares. El reto es acomodar profesores, grupos y materias sin violar reglas estrictas.",
    )
    add_note(
        doc,
        "Explicacion tecnica",
        "Uso un algoritmo genetico. Cada individuo es un horario completo representado como un vector de bloques. El algoritmo evalua cada horario con una funcion fitness, penalizando choques de profesor, choques de grupo, horas sin asignar y huecos.",
    )
    add_note(
        doc,
        "Cierre",
        "La aplicacion no solo genera un horario, tambien permite validar si la solucion es buena mediante metricas. El resultado ideal es que las reglas rotas sean cero y que la grafica de convergencia muestre mejora en el fitness.",
    )

    add_heading(doc, "10. Preguntas que podria hacer el profesor", 1)
    add_table(
        doc,
        ["Pregunta", "Respuesta sugerida"],
        [
            ["Que estructura usa el algoritmo?", "Usa un vector: cada posicion representa una sesion y cada valor indica el bloque de dia-hora asignado."],
            ["Por que es evolutivo?", "Porque genera una poblacion de soluciones, selecciona las mejores, cruza, muta y repite por generaciones."],
            ["Como evita choques?", "Los detecta en la funcion fitness y ademas aplica reparacion moviendo sesiones a bloques disponibles."],
            ["Que pasa si el fitness no es cero?", "No necesariamente significa que el horario sea invalido; hay que revisar primero que choques y horas sin asignar esten en cero."],
            ["Por que importar PDF?", "Porque los horarios reales suelen venir como cuadriculas; el sistema extrae esa informacion y la convierte en datos utilizables."],
        ],
        [2.05, 4.45],
    )

    add_heading(doc, "Conclusion", 1)
    doc.add_paragraph(
        "El proyecto demuestra la aplicacion de un algoritmo genetico a un problema escolar real. "
        "La solucion integra carga de datos, configuracion de restricciones, optimizacion evolutiva, reparacion de errores y visualizacion de resultados. "
        "La parte mas importante para defenderlo es explicar que el sistema no solo acomoda clases, sino que valida la calidad del horario mediante reglas estrictas y metricas."
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
